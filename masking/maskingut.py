import re
import pdfplumber
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PyPDF2 import PdfReader, PdfWriter
import os
from docx import Document
import uuid


def mask_text(content):
    """Mask email addresses and phone numbers."""
    email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b"
    phone_pattern = r"\b\d{10}\b|\b(?:\d{3}-){2}\d{4}\b|\b\(\d{3}\) \d{3}-\d{4}\b"
    masked_content = re.sub(email_pattern, "****@****.***", content)
    masked_content = re.sub(phone_pattern, "***-***-****", masked_content)
    return masked_content


def mask_pdf(input_path, output_path):
    """Mask content in a PDF file."""
    overlay_path = f"temp_{uuid.uuid4()}.pdf"

    with pdfplumber.open(input_path) as pdf:
        first_page = pdf.pages[0]
        width = float(first_page.mediabox[2])
        height = float(first_page.mediabox[3])

        c = canvas.Canvas(overlay_path, pagesize=(width, height))

        for page_num, page in enumerate(pdf.pages):
            text = page.extract_text()
            if not text:
                continue

            words = page.extract_words()
            for word in words:
                if re.search(
                    r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b",
                    word["text"],
                ):
                    x, y = word["x0"], height - word["top"]
                    w = word["x1"] - word["x0"]
                    h = word["bottom"] - word["top"]
                    c.setFillColorRGB(1, 1, 1)
                    c.rect(x, y - h, w, h, fill=True, stroke=False)
                    c.setFillColorRGB(0, 0, 0)
                    mask = "*" * (len(word["text"]) - 4)
                    c.drawString(x, y - h / 2, mask)

                elif re.search(
                    r"\b\d{10}\b|\b(?:\d{3}-){2}\d{4}\b|\b\(\d{3}\) \d{3}-\d{4}\b",
                    word["text"],
                ):
                    x, y = word["x0"], height - word["top"]
                    w = word["x1"] - word["x0"]
                    h = word["bottom"] - word["top"]
                    c.setFillColorRGB(1, 1, 1)
                    c.rect(x, y - h, w, h, fill=True, stroke=False)
                    c.setFillColorRGB(0, 0, 0)
                    mask = "*" * 10
                    c.drawString(x, y - h / 2, mask)

            c.showPage()
        c.save()

    reader = PdfReader(input_path)
    writer = PdfWriter()
    overlay_reader = PdfReader(overlay_path)

    for page_num, (orig_page, overlay_page) in enumerate(
        zip(reader.pages, overlay_reader.pages)
    ):
        orig_page.merge_page(overlay_page)
        writer.add_page(orig_page)

    with open(output_path, "wb") as f:
        writer.write(f)

    os.remove(overlay_path)  # Cleanup temporary file


def mask_docx(input_path, output_path):
    """Mask content in a DOCX file."""
    doc = Document(input_path)
    for paragraph in doc.paragraphs:
        paragraph.text = mask_text(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = mask_text(paragraph.text)
    doc.save(output_path)


def mask_txt(input_path, output_path):
    """Mask content in a TXT file."""
    with open(input_path, "r") as f:
        content = f.read()
    masked_content = mask_text(content)
    with open(output_path, "w") as f:
        f.write(masked_content)


# FastAPI setup
